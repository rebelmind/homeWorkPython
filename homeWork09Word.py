from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document()
doc.add_paragraph("Hello Python")
doc.save("hello_python.docx")

doc = Document("hello_python.docx")
bold_text = ""

for para in doc.paragraphs:
    for run in para.runs:
        if run.bold:  
            bold_text += run.text

print("Жирный текст:", bold_text)

new_doc = Document()
para = new_doc.add_paragraph("new paragraph")

run = para.runs[0]
run.font.name = "Calibri"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
run.font.size = Pt(14)
new_doc.save("newWord.docx")
